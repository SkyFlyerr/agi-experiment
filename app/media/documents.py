"""Document processing (PDF, DOCX, TXT) for text extraction."""

import logging
from typing import Dict, Any

from .utils import get_file_extension, validate_file_size

logger = logging.getLogger(__name__)

# Supported document formats
SUPPORTED_DOCUMENT_FORMATS = {".pdf", ".docx", ".txt"}


async def process_document(
    file_path: str,
    mime_type: str = "",
) -> Dict[str, Any]:
    """
    Extract text from document.

    Args:
        file_path: Path to document file
        mime_type: MIME type (optional, will be detected from extension)

    Returns:
        Dictionary with extraction result:
        {
            'status': 'success' | 'error',
            'text': str (if success),
            'page_count': int (if PDF),
            'word_count': int,
            'format': str,
            'error': str (if error),
        }
    """
    try:
        # Validate file size (10MB limit)
        if not validate_file_size(file_path, 10):
            return {
                "status": "error",
                "error": "File size exceeds 10MB limit",
            }

        # Validate format
        ext = get_file_extension(file_path)
        if ext not in SUPPORTED_DOCUMENT_FORMATS:
            return {
                "status": "error",
                "error": f"Unsupported format: {ext}",
                "supported_formats": list(SUPPORTED_DOCUMENT_FORMATS),
            }

        # Route to specific handler
        if ext == ".pdf":
            return await _extract_from_pdf(file_path)
        elif ext == ".docx":
            return await _extract_from_docx(file_path)
        elif ext == ".txt":
            return await _extract_from_txt(file_path)
        else:
            return {
                "status": "error",
                "error": f"No handler for format: {ext}",
            }

    except Exception as e:
        logger.error(f"Error processing document: {e}")
        return {
            "status": "error",
            "error": str(e),
        }


async def _extract_from_pdf(file_path: str) -> Dict[str, Any]:
    """Extract text from PDF."""
    try:
        import pypdf

        text_content = []
        page_count = 0

        with open(file_path, "rb") as f:
            pdf_reader = pypdf.PdfReader(f)
            page_count = len(pdf_reader.pages)

            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)

        full_text = "\n".join(text_content)
        word_count = len(full_text.split())

        result = {
            "status": "success",
            "text": full_text,
            "page_count": page_count,
            "word_count": word_count,
            "format": ".pdf",
        }

        logger.info(f"Extracted text from PDF ({page_count} pages, {word_count} words)")
        return result

    except ImportError:
        logger.error("pypdf not installed. Install with: pip install pypdf")
        return {
            "status": "error",
            "error": "pypdf not installed",
        }

    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return {
            "status": "error",
            "error": f"PDF extraction failed: {str(e)}",
        }


async def _extract_from_docx(file_path: str) -> Dict[str, Any]:
    """Extract text from DOCX."""
    try:
        from docx import Document

        doc = Document(file_path)
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)

        full_text = "\n".join(text_content)
        word_count = len(full_text.split())

        result = {
            "status": "success",
            "text": full_text,
            "word_count": word_count,
            "format": ".docx",
        }

        logger.info(f"Extracted text from DOCX ({word_count} words)")
        return result

    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return {
            "status": "error",
            "error": "python-docx not installed",
        }

    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return {
            "status": "error",
            "error": f"DOCX extraction failed: {str(e)}",
        }


async def _extract_from_txt(file_path: str) -> Dict[str, Any]:
    """Extract text from TXT."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        word_count = len(text.split())

        result = {
            "status": "success",
            "text": text,
            "word_count": word_count,
            "format": ".txt",
        }

        logger.info(f"Extracted text from TXT ({word_count} words)")
        return result

    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()

            word_count = len(text.split())

            return {
                "status": "success",
                "text": text,
                "word_count": word_count,
                "format": ".txt",
            }

        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return {
                "status": "error",
                "error": f"TXT extraction failed: {str(e)}",
            }

    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return {
            "status": "error",
            "error": f"TXT extraction failed: {str(e)}",
        }


__all__ = ["process_document"]
